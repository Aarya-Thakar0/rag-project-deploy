"""
Flask web application for document ingestion pipeline.
"""
''' 
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from groq import Groq
import os
import time
from pathlib import Path
from dotenv import load_dotenv
import docx
import numpy as np
from sklearn.feature_extraction.text import HashingVectorizer

app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()

# Global vector store
vector_store = None
chat_history = []
summaries = {}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'data', 'raw_docs')

def cleanup_old_documents():
    docs_folder = app.config['UPLOAD_FOLDER']
    if os.path.exists(docs_folder):
        try:
            for file in os.listdir(docs_folder):
                if file.lower().endswith(('.pdf', '.docx', '.txt')):
                    filepath = os.path.join(docs_folder, file)
                    os.remove(filepath)
        except Exception:
            pass

# Clean up old documents on startup
cleanup_old_documents()

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_document(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == 'docx':
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return [Document(page_content='\n'.join(full_text), metadata={"source": file_path})]
    elif ext == 'txt':
        loader = TextLoader(file_path)
        return loader.load()
    return []

class LocalHashEmbeddings:
    def __init__(self, n_features: int = 2048):
        self.vectorizer = HashingVectorizer(
            n_features=n_features,
            alternate_sign=False,
            norm='l2',
            lowercase=True,
            stop_words='english'
        )
    def embed_documents(self, texts):
        X = self.vectorizer.transform(texts).toarray()
        return [x.tolist() for x in X]
    def embed_query(self, text):
        x = self.vectorizer.transform([text]).toarray()[0]
        return x.tolist()
    def __call__(self, text_or_texts):
        if isinstance(text_or_texts, (list, tuple)):
            return self.embed_documents(list(text_or_texts))
        return self.embed_query(text_or_texts)

def summarize_text(text):
    if not GROQ_API_KEY:
        return "Summary not available: GROQ_API_KEY not set."
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model="grok-1",
            messages=[
                {"role": "system", "content": "Provide a concise summary of the following text in 2-3 sentences."},
                {"role": "user", "content": text[:8000]}  # Limit to 8000 chars
            ],
            temperature=0.3,
            max_completion_tokens=1024,
            top_p=1,
        )
        summary = response.choices[0].message.content.strip()
        if not summary:
            return "Summary could not be generated."
        return summary
    except Exception as e:
        return f"Summary error: {str(e)}"

def process_document(file_path):
    global vector_store
    try:
        raw_docs = load_document(file_path)
        if not raw_docs:
            return {'success': False, 'error': 'Document has no extractable text content.'}
        total_chars = sum(len(doc.page_content.strip()) for doc in raw_docs)
        if total_chars == 0:
            return {'success': False, 'error': 'Document contains no text. It might be empty or scanned.'}
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_documents(raw_docs)
        vector_store = None
        batch_size = 5
        embeddings = LocalHashEmbeddings(n_features=2048)
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            if vector_store is None:
                vector_store = FAISS.from_documents(batch, embeddings)
            else:
                vector_store.add_documents(batch)
            time.sleep(1)
        # Generate summary
        full_text = "\n".join([doc.page_content for doc in raw_docs])
        filename = os.path.basename(file_path)
        summaries[filename] = summarize_text(full_text)
        return {
            'success': True,
            'document_count': len(raw_docs),
            'chunk_count': len(chunks),
            'chunks': [chunk.page_content[:200] + "..." for chunk in chunks[:5]]
        }
    except Exception as e:
        error_msg = f"Error processing document: {str(e)}"
        return {'success': False, 'error': error_msg}

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/api/process', methods=['POST'])
def api_process():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF, DOCX, and TXT files are allowed'}), 400
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    result = process_document(filepath)
    return jsonify(result)

@app.route('/api/documents')
def api_documents():
    docs_folder = app.config['UPLOAD_FOLDER']
    documents = []
    os.makedirs(docs_folder, exist_ok=True)
    if os.path.exists(docs_folder):
        try:
            for file in os.listdir(docs_folder):
                if file.lower().endswith(('.pdf', '.docx', '.txt')):
                    filepath = os.path.join(docs_folder, file)
                    size = os.path.getsize(filepath)
                    documents.append({'name': file, 'size': f"{size / (1024*1024):.2f} MB"})
        except Exception:
            pass
    return jsonify({'documents': documents})

@app.route('/api/summaries')
def api_summaries():
    return jsonify({'summaries': summaries})

@app.route('/api/summary/<filename>')
def api_summary(filename):
    filename = secure_filename(filename)
    if filename in summaries:
        return jsonify({'summary': summaries[filename]})
    return jsonify({'error': 'Summary not available'}), 404

@app.route('/api/process-existing/<filename>')
def api_process_existing(filename):
    filename = secure_filename(filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    result = process_document(filepath)
    return jsonify(result)

@app.route('/api/chat', methods=['POST'])
def api_chat():
    global vector_store, chat_history
    if not vector_store:
        return jsonify({'error': 'No document processed yet. Please upload a file first.'}), 400
    data = request.json
    question = data.get('question')
    if not question:
        return jsonify({'error': 'No question provided'}), 400
    try:
        k = 4
        source_docs = vector_store.similarity_search(question, k=k)
        context_parts = []
        for doc in source_docs:
            content = doc.page_content
            if len(content) > 1200:
                content = content[:1200] + "..."
            context_parts.append(content)
        context_text = "\n\n".join(context_parts)
        if not GROQ_API_KEY:
            return jsonify({'error': 'GROQ_API_KEY not set. Please add it to .env and restart.'}), 401
        client = Groq(api_key=GROQ_API_KEY)
        completion = client.chat.completions.create(
            model="grok-1",
            messages=[
                {"role": "system", "content": "Answer based only on the provided context. If missing, say you lack information."},
                {"role": "user", "content": f"Context:\n{context_text}\n\nQuestion:\n{question}"}
            ],
            temperature=1,
            max_completion_tokens=8192,
            top_p=1,
            reasoning_effort="medium",
            stream=True,
            stop=None
        )
        answer_text = ""
        for chunk in completion:
            delta = chunk.choices[0].delta
            if delta and hasattr(delta, "content") and delta.content:
                answer_text += delta.content
        sources = [doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content for doc in source_docs]
        chat_history.append({"role": "user", "content": question})
        chat_history.append({"role": "assistant", "content": answer_text})
        return jsonify({"answer": answer_text, "sources": sources})
    except Exception as e:
        err_str = str(e)
        if 'invalid_api_key' in err_str.lower():
            return jsonify({'error': 'Invalid GROQ_API_KEY. Please verify the key in .env and restart.'}), 401
        return jsonify({'error': f"Error generating answer: {err_str}"}), 500

@app.route('/api/clear_chat', methods=['POST'])
def clear_chat():
    global chat_history
    chat_history.clear()
    return jsonify({'success': True, 'message': 'Chat history cleared'})

@app.errorhandler(404)
def not_found(error):
    if request.path.startswith('/api/'):
        return jsonify({'error': 'API endpoint not found'}), 404
    return '', 204

if __name__ == '__main__':
    app.run(debug=True, port=5000)


'''

"""
Flask web application for document ingestion pipeline.
"""
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from groq import Groq
import os
import time
from pathlib import Path
from dotenv import load_dotenv
import docx
import numpy as np
from sklearn.feature_extraction.text import HashingVectorizer

app = Flask(__name__, static_folder='static', static_url_path='/static')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()

# Global state
vector_store = None
chat_history = []
summaries = {}

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'data', 'raw_docs')

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

# ── Groq model to use ──────────────────────────────────────────────────────────
# "llama-3.3-70b-versatile" is a fast, capable model available on Groq.
# Change this to any model string from https://console.groq.com/docs/models
GROQ_MODEL = "llama-3.3-70b-versatile"


def cleanup_old_documents():
    docs_folder = app.config['UPLOAD_FOLDER']
    if os.path.exists(docs_folder):
        try:
            for file in os.listdir(docs_folder):
                if file.lower().endswith(('.pdf', '.docx', '.txt')):
                    filepath = os.path.join(docs_folder, file)
                    os.remove(filepath)
        except Exception:
            pass

# Clean up old docs on startup
cleanup_old_documents()


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def load_document(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'pdf':
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == 'docx':
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        return [Document(page_content='\n'.join(full_text), metadata={"source": file_path})]
    elif ext == 'txt':
        loader = TextLoader(file_path, encoding='utf-8')
        return loader.load()
    return []


class LocalHashEmbeddings:
    """Lightweight local embeddings using a HashingVectorizer — no API key needed."""

    def __init__(self, n_features: int = 2048):
        self.vectorizer = HashingVectorizer(
            n_features=n_features,
            alternate_sign=False,
            norm='l2',
            lowercase=True,
            stop_words='english'
        )

    def embed_documents(self, texts):
        X = self.vectorizer.transform(texts).toarray()
        return [x.tolist() for x in X]

    def embed_query(self, text):
        x = self.vectorizer.transform([text]).toarray()[0]
        return x.tolist()

    def __call__(self, text_or_texts):
        if isinstance(text_or_texts, (list, tuple)):
            return self.embed_documents(list(text_or_texts))
        return self.embed_query(text_or_texts)


def summarize_text(text: str) -> str:
    """Generate a 2–3 sentence summary using Groq."""
    if not GROQ_API_KEY:
        return "Summary not available: GROQ_API_KEY not configured."
    try:
        client = Groq(api_key=GROQ_API_KEY)
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a concise document summarizer. "
                        "Summarize the provided text in 2–3 clear, informative sentences. "
                        "Focus on the main topic, key findings, and purpose of the document."
                    )
                },
                {"role": "user", "content": text[:8000]}
            ],
            temperature=0.3,
            max_tokens=300,
        )
        summary = response.choices[0].message.content.strip()
        return summary if summary else "Summary could not be generated."
    except Exception as e:
        return f"Summary error: {str(e)}"


def process_document(file_path: str) -> dict:
    """Chunk a document, build/update the FAISS vector store, and generate a summary."""
    global vector_store

    try:
        raw_docs = load_document(file_path)
        if not raw_docs:
            return {'success': False, 'error': 'Document has no extractable text content.'}

        total_chars = sum(len(doc.page_content.strip()) for doc in raw_docs)
        if total_chars == 0:
            return {'success': False, 'error': 'Document contains no text. It might be empty or scanned.'}

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_documents(raw_docs)

        vector_store = None
        embeddings = LocalHashEmbeddings(n_features=2048)
        batch_size = 5

        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            if vector_store is None:
                vector_store = FAISS.from_documents(batch, embeddings)
            else:
                vector_store.add_documents(batch)
            time.sleep(0.5)

        # Generate and cache the summary
        full_text = "\n".join(doc.page_content for doc in raw_docs)
        filename = os.path.basename(file_path)
        summaries[filename] = summarize_text(full_text)

        return {
            'success': True,
            'document_count': len(raw_docs),
            'chunk_count': len(chunks),
            'chunks': [chunk.page_content[:200] + "..." for chunk in chunks[:5]]
        }

    except Exception as e:
        return {'success': False, 'error': f"Error processing document: {str(e)}"}


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/favicon.ico')
def favicon():
    return '', 204


@app.route('/api/process', methods=['POST'])
def api_process():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF, DOCX, and TXT files are allowed'}), 400

    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    result = process_document(filepath)
    return jsonify(result)


@app.route('/api/documents')
def api_documents():
    docs_folder = app.config['UPLOAD_FOLDER']
    documents = []
    os.makedirs(docs_folder, exist_ok=True)

    try:
        for file in os.listdir(docs_folder):
            if file.lower().endswith(('.pdf', '.docx', '.txt')):
                filepath = os.path.join(docs_folder, file)
                size = os.path.getsize(filepath)
                documents.append({
                    'name': file,
                    'size': f"{size / (1024 * 1024):.2f} MB"
                })
    except Exception:
        pass

    return jsonify({'documents': documents})


@app.route('/api/summaries')
def api_summaries():
    return jsonify({'summaries': summaries})


@app.route('/api/summary/<filename>')
def api_summary(filename):
    filename = secure_filename(filename)

    # Return cached summary if available
    if filename in summaries:
        return jsonify({'summary': summaries[filename]})

    # If not cached but file exists, generate on the fly
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(filepath):
        try:
            raw_docs = load_document(filepath)
            if raw_docs:
                full_text = "\n".join(doc.page_content for doc in raw_docs)
                summary = summarize_text(full_text)
                summaries[filename] = summary
                return jsonify({'summary': summary})
        except Exception as e:
            return jsonify({'error': f'Could not generate summary: {str(e)}'}), 500

    return jsonify({'error': 'Summary not available and file not found'}), 404


@app.route('/api/process-existing/<filename>')
def api_process_existing(filename):
    filename = secure_filename(filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    result = process_document(filepath)
    return jsonify(result)


@app.route('/api/delete/<filename>', methods=['DELETE'])
def api_delete(filename):
    """Delete a specific uploaded document."""
    filename = secure_filename(filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404

    try:
        os.remove(filepath)
        # Remove from summaries cache
        summaries.pop(filename, None)
        return jsonify({'success': True, 'message': f'{filename} deleted successfully'})
    except Exception as e:
        return jsonify({'error': f'Could not delete file: {str(e)}'}), 500


@app.route('/api/chat', methods=['POST'])
def api_chat():
    global vector_store, chat_history

    if not vector_store:
        return jsonify({'error': 'No document processed yet. Please upload a file first.'}), 400

    data = request.json
    question = data.get('question', '').strip()

    if not question:
        return jsonify({'error': 'No question provided'}), 400

    if not GROQ_API_KEY:
        return jsonify({'error': 'GROQ_API_KEY not set. Please add it to .env and restart.'}), 401

    try:
        source_docs = vector_store.similarity_search(question, k=4)

        context_parts = []
        for doc in source_docs:
            content = doc.page_content
            if len(content) > 1200:
                content = content[:1200] + "..."
            context_parts.append(content)

        context_text = "\n\n".join(context_parts)

        client = Groq(api_key=GROQ_API_KEY)
        completion = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful AI assistant. Answer the user's question based ONLY on "
                        "the provided document context. If the context does not contain enough "
                        "information, say so clearly. Be concise and accurate."
                    )
                },
                {
                    "role": "user",
                    "content": f"Context:\n{context_text}\n\nQuestion:\n{question}"
                }
            ],
            temperature=0.7,
            max_tokens=2048,
            stream=True,
        )

        answer_text = ""
        for chunk in completion:
            delta = chunk.choices[0].delta
            if delta and hasattr(delta, "content") and delta.content:
                answer_text += delta.content

        sources = [
            doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content
            for doc in source_docs
        ]

        chat_history.append({"role": "user", "content": question})
        chat_history.append({"role": "assistant", "content": answer_text})

        return jsonify({"answer": answer_text, "sources": sources})

    except Exception as e:
        err_str = str(e)
        if 'invalid_api_key' in err_str.lower() or 'authentication' in err_str.lower():
            return jsonify({'error': 'Invalid GROQ_API_KEY. Please verify the key in .env.'}), 401
        if 'model_not_found' in err_str.lower() or 'does not exist' in err_str.lower():
            return jsonify({'error': f'Model "{GROQ_MODEL}" not found. Check available models at console.groq.com.'}), 500
        return jsonify({'error': f"Error generating answer: {err_str}"}), 500


@app.route('/api/clear_chat', methods=['POST'])
def clear_chat():
    global chat_history
    chat_history.clear()
    return jsonify({'success': True, 'message': 'Chat history cleared'})


# ── Error handlers ────────────────────────────────────────────────────────────

@app.errorhandler(404)
def not_found(error):
    if request.path.startswith('/api/'):
        return jsonify({'error': 'API endpoint not found'}), 404
    return '', 204


@app.errorhandler(413)
def too_large(error):
    return jsonify({'error': 'File too large. Maximum size is 50MB.'}), 413


if __name__ == '__main__':
    app.run(debug=True, port=5000)